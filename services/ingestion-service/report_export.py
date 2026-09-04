"""
Render a stored Markdown report into a downloadable DOCX or PDF file.

Both renderers are pure-Python (python-docx, fpdf2) so the container needs no
system packages. The Markdown handled here is the small subset the report
worker produces: '#'/'##'/'###' headings, '**bold**', '- ' / '* ' bullets,
and plain paragraphs.
"""
import io
import re
from datetime import datetime, timezone
from typing import Optional, List, Dict, Any


def _clean_inline(text: str) -> str:
    """Strip Markdown emphasis markers for renderers that style separately."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'(?<!\*)\*(?!\*)(.+?)\*(?!\*)', r'\1', text)
    text = text.replace('`', '')
    return text.strip()


def _parse_blocks(md: str) -> List[Dict[str, Any]]:
    """Turn Markdown text into a flat list of typed blocks."""
    blocks: List[Dict[str, Any]] = []
    for raw in (md or "").splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        h = re.match(r'^(#{1,6})\s+(.*)$', line)
        if h:
            blocks.append({"type": "heading", "level": len(h.group(1)), "text": _clean_inline(h.group(2))})
            continue
        b = re.match(r'^\s*[-*]\s+(.*)$', line)
        if b:
            blocks.append({"type": "bullet", "text": _clean_inline(b.group(1))})
            continue
        blocks.append({"type": "para", "text": _clean_inline(line)})
    return blocks


def _sources_lines(sources: Optional[List[Dict[str, Any]]]) -> List[str]:
    out = []
    for i, s in enumerate(sources or [], start=1):
        name = s.get("filename") or s.get("document_id") or "source"
        sub = s.get("subsidiary")
        page = s.get("page_number") or s.get("chunk_index")
        extra = " ".join(x for x in [f"({sub})" if sub else "", f"chunk #{page}" if page is not None else ""] if x)
        out.append(f"{i}. {name} {extra}".strip())
    return out


def render_docx(report: Dict[str, Any], document: Dict[str, Any]) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    run = title.add_run("MineIQ — Automated Report")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)

    meta = doc.add_paragraph()
    meta_run = meta.add_run(
        f"{document.get('original_filename', '')}   ·   "
        f"{document.get('subsidiary') or 'N/A'}   ·   {document.get('doc_type') or 'unclassified'}"
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    stamp = doc.add_paragraph()
    stamp_run = stamp.add_run(
        f"Generated {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}   ·   AI-GENERATED DRAFT — requires human review"
    )
    stamp_run.font.size = Pt(8)
    stamp_run.italic = True
    stamp_run.font.color.rgb = RGBColor(0x9A, 0x3B, 0x3B)

    doc.add_paragraph()

    for blk in _parse_blocks(report.get("report_text", "")):
        if blk["type"] == "heading":
            doc.add_heading(blk["text"], level=min(blk["level"], 4))
        elif blk["type"] == "bullet":
            doc.add_paragraph(blk["text"], style="List Bullet")
        else:
            doc.add_paragraph(blk["text"])

    src_lines = _sources_lines(report.get("sources"))
    if src_lines:
        doc.add_heading("Source references", level=2)
        for line in src_lines:
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_pdf(report: Dict[str, Any], document: Dict[str, Any]) -> bytes:
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    MARGIN = 18
    pdf = FPDF(format="A4")
    pdf.set_auto_page_break(auto=True, margin=MARGIN)
    pdf.set_margins(MARGIN, MARGIN, MARGIN)
    pdf.add_page()
    width = pdf.w - 2 * MARGIN  # explicit content width; multi_cell(0,..) can be 0 here

    def _txt(s: str) -> str:
        # fpdf core fonts are latin-1 only; drop anything outside it.
        return (s or "").encode("latin-1", "replace").decode("latin-1")

    def cell(h: float, text: str):
        pdf.multi_cell(width, h, _txt(text), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.set_text_color(31, 42, 68)
    pdf.set_font("Helvetica", "B", 20)
    cell(10, "MineIQ - Automated Report")

    pdf.set_text_color(107, 114, 128)
    pdf.set_font("Helvetica", "", 9)
    cell(5, f"{document.get('original_filename', '')}   -   "
            f"{document.get('subsidiary') or 'N/A'}   -   {document.get('doc_type') or 'unclassified'}")
    pdf.set_text_color(154, 59, 59)
    pdf.set_font("Helvetica", "I", 8)
    cell(5, f"Generated {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}   -   "
            "AI-GENERATED DRAFT - requires human review")
    pdf.ln(3)

    for blk in _parse_blocks(report.get("report_text", "")):
        if blk["type"] == "heading":
            size = {1: 15, 2: 13, 3: 12}.get(blk["level"], 11)
            pdf.set_text_color(31, 42, 68)
            pdf.set_font("Helvetica", "B", size)
            pdf.ln(2)
            cell(7, blk["text"])
        elif blk["type"] == "bullet":
            pdf.set_text_color(30, 30, 30)
            pdf.set_font("Helvetica", "", 11)
            cell(6, f"  -  {blk['text']}")
        else:
            pdf.set_text_color(30, 30, 30)
            pdf.set_font("Helvetica", "", 11)
            cell(6, blk["text"])

    src_lines = _sources_lines(report.get("sources"))
    if src_lines:
        pdf.ln(2)
        pdf.set_text_color(31, 42, 68)
        pdf.set_font("Helvetica", "B", 13)
        cell(7, "Source references")
        pdf.set_text_color(30, 30, 30)
        pdf.set_font("Helvetica", "", 10)
        for line in src_lines:
            cell(6, line)

    out = pdf.output()
    return bytes(out)
